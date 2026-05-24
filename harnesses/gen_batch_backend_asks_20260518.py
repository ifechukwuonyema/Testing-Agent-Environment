"""
Generates batch_backend_asks_20260518.docx
Run 8 baseline: 85P/110F/27B = 38.3% (222 TCs)
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = r"C:\Users\Onyema Ifechukwu\Downloads\Kardit\reports\batch_backend_asks_20260518.docx"
HDR_COLOR = RGBColor(0x1F, 0x3A, 0x6E)
HDR_TXT   = RGBColor(0xFF, 0xFF, 0xFF)
RED       = RGBColor(0xC0, 0x00, 0x00)
ORANGE    = RGBColor(0xED, 0x7D, 0x31)
YELLOW    = RGBColor(0xFF, 0xC0, 0x00)
GREEN     = RGBColor(0x70, 0xAD, 0x47)
GRAY      = RGBColor(0xBF, 0xBF, 0xBF)


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = str(rgb)  # RGBColor.__str__ returns 6-char uppercase hex
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def hdr_row(table, cols):
    row = table.rows[0]
    for i, text in enumerate(cols):
        cell = row.cells[i]
        cell.text = text
        set_cell_bg(cell, HDR_COLOR)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = HDR_TXT
        run.font.size = Pt(9)


def add_table_row(table, values, bold_first=False, bg: RGBColor = None):
    row = table.add_row()
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        p = cell.paragraphs[0]
        run = p.runs[0] if p.runs else p.add_run(str(val))
        run.font.size = Pt(9)
        if bold_first and i == 0:
            run.bold = True
        if bg:
            set_cell_bg(cell, bg)
    return row


def style_heading(para, level=1):
    para.style = f"Heading {level}"


doc = Document()

# --- margins ---
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# ═══════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════
title = doc.add_heading("Kardit Batch Microservice", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading("Backend Recommendations & Provisions Report", 1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta_pairs = [
    ("Service",        "Batch"),
    ("Run Reference",  "Run 8 — 2026-05-18"),
    ("Test Result",    "85P / 110F / 27B = 38.3%  (222 TCs, 7 endpoints)"),
    ("Tester",         "postman_hybrid_batch_runner"),
    ("Environment",    "http://167.172.49.177:8080"),
    ("Affiliate Seed", "AFF-9F6EDBBE20DD4C6B97D0B720676506E1"),
    ("Processing ID",  "952480b6-61d2-4299-a6ca-430dce7a316c"),
    ("Completed ID",   "ef57c562-4a98-4c46-b8ec-13e36a1a3ebe"),
    ("Download Token", "6483757b22984a2884995a9ec1a6fe10"),
    ("Report Date",    "2026-05-18"),
]
for k, v in meta_pairs:
    p = doc.add_paragraph()
    r = p.add_run(f"{k}:  ")
    r.bold = True
    r.font.size = Pt(10)
    r2 = p.add_run(v)
    r2.font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════
doc.add_heading("1. Executive Summary", 1)

doc.add_paragraph(
    "Run 8 exercised all 7 Batch API endpoints across 222 test cases using the "
    "postman_hybrid_batch_runner. The overall pass rate is 38.3% (85/222). "
    "27 test cases are structurally BLOCKED and cannot be unblocked by the harness alone "
    "(15 Cluster-C persistence splits, 6 DB-verify-only, 5 rate-flood, 1 unrecognised scenario). "
    "Of the 195 executable test cases, 85 pass (43.6%)."
)

doc.add_paragraph(
    "Three backend defects each individually prevent entire endpoint clusters from reaching "
    "a passing state: (1) the token-download handler (BATCH-07) is completely unimplemented "
    "or crashes unconditionally; (2) the batch state machine auto-advances UPLOADED→PROCESSING "
    "before validate or submit can fire; and (3) authentication/authorisation middleware is not "
    "applied to any Batch endpoint. Resolving these three defects alone lifts the ceiling from "
    "43.6% to an estimated ≥ 81%."
)

# Per-endpoint table
doc.add_heading("Endpoint Pass-Rate Summary", 2)

ep_tbl = doc.add_table(rows=1, cols=6)
ep_tbl.style = "Table Grid"
ep_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(ep_tbl, ["Endpoint", "API ID", "PASS", "FAIL", "BLOCKED", "Rate"])

endpoint_rows = [
    ("POST /api/v1/Batches/card-creation/upload",         "BATCH-01", 16, 12,  2, "53.3%"),
    ("POST /api/v1/Batches/{batchId}/validate",           "BATCH-02",  3, 25,  2, "10.0%"),
    ("POST /api/v1/Batches/{batchId}/submit",             "BATCH-03",  9, 23,  2, "26.5%"),
    ("GET /api/v1/Batches/{batchId}",                     "BATCH-04", 30,  4,  0, "88.2%"),
    ("GET /api/v1/Batches/{batchId}/rows",                "BATCH-05",  7,  8, 15, "46.7%"),
    ("GET /api/v1/Batches/{batchId}/results/download",    "BATCH-06", 19, 11,  2, "59.4%"),
    ("GET /api/v1/Batches/{batchId}/results/download/{token}", "BATCH-07",  1, 27,  4, " 3.6%"),
]
for ep, aid, p, f, b, rate in endpoint_rows:
    add_table_row(ep_tbl, [ep, aid, p, f, b, rate])

doc.add_paragraph()

# ═══════════════════════════════════════════════════════
# 2. DEFECT PRIORITY MATRIX
# ═══════════════════════════════════════════════════════
doc.add_heading("2. Defect Priority Matrix", 1)

pri_tbl = doc.add_table(rows=1, cols=6)
pri_tbl.style = "Table Grid"
pri_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(pri_tbl, ["Defect ID", "Endpoint(s)", "Severity", "FAIL Impact", "Root Cause", "Fix Difficulty"])

defects = [
    ("D-BATCH-07-CRASH",  "BATCH-07",                    "Critical", "+26F",  "Unimplemented/crashing token handler",      "Medium"),
    ("D-BATCH-AUTH-1",    "BATCH-01/02/03/04/06/07",     "Critical", "+35F",  "Auth middleware absent on all endpoints",   "Low"),
    ("D-BATCH-02-STATE",  "BATCH-02",                    "High",     "+22F",  "UPLOADED→PROCESSING auto-transition race",  "Medium"),
    ("D-BATCH-03-STATE",  "BATCH-03",                    "High",     "+17F",  "Submit requires VALIDATED state (never set)","Medium"),
    ("D-BATCH-05-ROWS",   "BATCH-05",                    "High",     "+15B",  "Rows not persisted / wrong ID scope",       "Medium"),
    ("D-BATCH-01-VAL",    "BATCH-01",                    "Medium",   "+8F",   "File-type/CSV-header validation absent",    "Low"),
    ("D-BATCH-01-PROD",   "BATCH-01",                    "Medium",   "+4F",   "ProductId/size constraint not enforced",    "Low"),
    ("D-BATCH-06-STATE",  "BATCH-06",                    "Medium",   "+6F",   "Download requires COMPLETED state batch",   "Medium"),
    ("D-BATCH-04-AUTH",   "BATCH-04",                    "Medium",   "+4F",   "Auth bypass on GET batch status endpoint",  "Low"),
]

sev_colors = {
    "Critical": RED,
    "High":     ORANGE,
    "Medium":   YELLOW,
    "Low":      GREEN,
}
for defect_id, endpoints, sev, impact, cause, fix_diff in defects:
    row = add_table_row(pri_tbl, [defect_id, endpoints, sev, impact, cause, fix_diff])
    for i, cell in enumerate(row.cells):
        if i == 2:
            set_cell_bg(cell, sev_colors.get(sev, GRAY))

doc.add_paragraph()

# ═══════════════════════════════════════════════════════
# 3. DETAILED DEFECT REPORTS
# ═══════════════════════════════════════════════════════
doc.add_heading("3. Detailed Defect Reports", 1)

# ── D-BATCH-07-CRASH ──────────────────────────────────
doc.add_heading("D-BATCH-07-CRASH — Token Download Handler: Unimplemented / Crash", 2)
fields = [
    ("Defect ID",        "D-BATCH-07-CRASH"),
    ("Severity",         "Critical"),
    ("Endpoint",         "GET /api/v1/Batches/{batchId}/results/download/{token}  (BATCH-07)"),
    ("Affected TCs",     "27 of 32 return 500 INTERNAL_SERVER_ERROR"),
    ("Sole PASS",        "malformed_batchid_rejected → 400 (batchId format validation fires before the crash)"),
    ("Root Cause",       "The token-download route handler is either unimplemented or throws an unhandled exception. "
                         "No matter what token value is supplied the backend crashes."),
    ("Evidence",         'HTTP 500: {"status":"error","error":{"code":"INTERNAL_SERVER_ERROR",'
                         '"message":"An unexpected error occurred. Request id: 0HNLKUCB44FUL:00000001"}}'),
    ("Observed behaviour","All token variations (valid, invalid, expired, oversized, special-char, nil-uuid) return 500. "
                           "Only a malformed batchId triggers a 400 before the handler is reached."),
    ("Expected behaviour","Valid token + matching batchId → 200 with file content (Content-Disposition: attachment). "
                           "Invalid/expired/mismatched token → 404 or 401/403 per scenario."),
    ("Recommendation",   "Implement the token-download route handler. At minimum, add a try/catch that returns 500 "
                         "with a structured error instead of an unhandled crash. Then enforce token validation: "
                         "existence check → expiry check → batchId ownership check → serve file."),
    ("Ceiling unlock",   "Fixes ~26 FAILs on BATCH-07 (+11.7 pp platform-wide)"),
]
for k, v in fields:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f"{k}: ")
    r.bold = True
    r.font.size = Pt(9)
    p.add_run(v).font.size = Pt(9)

doc.add_paragraph()

# ── D-BATCH-AUTH-1 ──────────────────────────────────
doc.add_heading("D-BATCH-AUTH-1 — Authentication / Authorisation Middleware Missing", 2)
fields = [
    ("Defect ID",        "D-BATCH-AUTH-1"),
    ("Severity",         "Critical"),
    ("Endpoints",        "All 7 Batch endpoints (BATCH-01 through BATCH-07)"),
    ("Affected TCs",     "≥ 35 FAILs across all endpoints; all auth-bypass and role-rejection scenarios"),
    ("Root Cause",       "No authentication or authorisation middleware is registered on any Batch route. "
                         "Requests without a token, with an invalid token, or using the wrong user role "
                         "all return 200 (success) instead of 401 / 403."),
    ("Evidence BATCH-01","bank_user_upload_rejected → expected 403, got 200 (batchId returned, status UPLOADED)\n"
                         "service_provider_upload_rejected → expected 403, got 200\n"
                         "unauthenticated_upload_rejected → expected 401, got 200\n"
                         "foreign_tenant_context_rejected → expected 403, got 200"),
    ("Evidence BATCH-04","unauthenticated_get_rejected → expected 401, got 200\n"
                         "foreign_scope_batch_rejected → expected 403, got 200"),
    ("Expected behaviour","Any request missing a valid bearer token → 401 Unauthorized.\n"
                         "Requests with a valid token but wrong role (Bank, ServiceProvider) → 403 Forbidden.\n"
                         "Requests using a different affiliate/tenant context → 403 Forbidden."),
    ("Recommendation",   "Register JWT/bearer auth middleware on all Batch routes. "
                         "Enforce role-based access control: only Affiliate users with matching tenantId/affiliateId "
                         "may create or manage batches. Bank and ServiceProvider roles must be explicitly rejected."),
    ("Ceiling unlock",   "Fixes ~35 FAILs across 6 endpoints (+15.8 pp platform-wide)"),
]
for k, v in fields:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f"{k}: ")
    r.bold = True
    r.font.size = Pt(9)
    p.add_run(v).font.size = Pt(9)

doc.add_paragraph()

# ── D-BATCH-02-STATE ──────────────────────────────────
doc.add_heading("D-BATCH-02-STATE — Validate Endpoint: UPLOADED→PROCESSING Auto-Transition Race", 2)
fields = [
    ("Defect ID",        "D-BATCH-02-STATE"),
    ("Severity",         "High"),
    ("Endpoint",         "POST /api/v1/Batches/{batchId}/validate  (BATCH-02)"),
    ("Affected TCs",     "22 FAILs on BATCH-02; only 3 PASSes (happy-path variants that hit the state check before auto-advance)"),
    ("Root Cause",       "A freshly uploaded batch transitions UPLOADED → PROCESSING automatically (triggered by the upload handler or a background job). "
                         "By the time the validate endpoint is called, the batch is already in PROCESSING state. "
                         "The validate handler correctly rejects non-UPLOADED batches, returning: "
                         "'Batch job must be in UPLOADED status before validation can begin' (HTTP 409). "
                         "This is not a validate defect — it is a state-machine design flaw: "
                         "validate should be callable before auto-processing begins, or auto-processing should be gated on explicit validate/submit."),
    ("Evidence",         "POST .../validate → HTTP 409: 'Batch job must be in UPLOADED status before validation can begin'\n"
                         "GET .../status immediately after upload → status: PROCESSING (not UPLOADED)\n"
                         "BATCH-01 upload response → status: UPLOADED (correct at upload time)\n"
                         "BATCH-04 GET → status: PROCESSING (auto-advance confirmed)"),
    ("Expected behaviour","After upload the batch stays in UPLOADED until the affiliate explicitly calls validate. "
                          "Validate transitions UPLOADED → VALIDATING → VALIDATED (or VALIDATION_FAILED). "
                          "Submit is only callable after VALIDATED."),
    ("Recommendation",   "Option A (preferred): Remove the automatic UPLOADED→PROCESSING transition. "
                         "Batches should remain in UPLOADED until validate is explicitly called.\n"
                         "Option B: If auto-processing must remain, add a 'quick-validate' call in the upload handler "
                         "before the state transition, and expose the result on the batch resource.\n"
                         "Option C (short-term): Provide the harness with a batch that is guaranteed to be in UPLOADED state "
                         "at test time (see Provisions section)."),
    ("Ceiling unlock",   "Fixes ~22 FAILs on BATCH-02 (+9.9 pp platform-wide)"),
]
for k, v in fields:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f"{k}: ")
    r.bold = True
    r.font.size = Pt(9)
    p.add_run(v).font.size = Pt(9)

doc.add_paragraph()

# ── D-BATCH-03-STATE ──────────────────────────────────
doc.add_heading("D-BATCH-03-STATE — Submit Endpoint: VALIDATED State Never Reachable", 2)
fields = [
    ("Defect ID",        "D-BATCH-03-STATE"),
    ("Severity",         "High"),
    ("Endpoint",         "POST /api/v1/Batches/{batchId}/submit  (BATCH-03)"),
    ("Affected TCs",     "23 FAILs; 9 PASSes (scenarios that test non-state-dependent paths)"),
    ("Root Cause",       "Submit requires VALIDATED state. Because validate is blocked by D-BATCH-02-STATE "
                         "(batch auto-transitions past UPLOADED before validate fires), no batch ever reaches "
                         "VALIDATED state during a test session. Submit therefore always rejects with a state error."),
    ("Evidence",         "POST .../submit → 409 or 422: batch is not in VALIDATED state"),
    ("Dependency",       "Blocked by D-BATCH-02-STATE. Fixing the validate flow unblocks submit."),
    ("Recommendation",   "Fix D-BATCH-02-STATE first. Then provide a pre-validated batch ID for testing "
                         "(see Provisions: VALIDATED_BATCH_ID)."),
    ("Ceiling unlock",   "Fixes ~17 FAILs on BATCH-03 (+7.7 pp platform-wide) once D-BATCH-02-STATE is resolved"),
]
for k, v in fields:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f"{k}: ")
    r.bold = True
    r.font.size = Pt(9)
    p.add_run(v).font.size = Pt(9)

doc.add_paragraph()

# ── D-BATCH-05-ROWS ──────────────────────────────────
doc.add_heading("D-BATCH-05-ROWS — Rows Endpoint: Cluster-C Persistence Split", 2)
fields = [
    ("Defect ID",        "D-BATCH-05-ROWS"),
    ("Severity",         "High"),
    ("Endpoint",         "GET /api/v1/Batches/{batchId}/rows  (BATCH-05)"),
    ("Affected TCs",     "15 BLOCKEDs (Cluster-C), 8 FAILs"),
    ("Root Cause",       "The same batchId returns HTTP 200 on GET /api/v1/Batches/{batchId} (BATCH-04) "
                         "but HTTP 404 on GET /api/v1/Batches/{batchId}/rows. "
                         "This indicates a write/read consistency defect: the batch record is persisted to the "
                         "primary batch table but the rows table/projection is either not populated or uses a "
                         "different key space."),
    ("Evidence",         "Seeded batchId 085179c3-cbd6-4844-908c-52dcb3c90929:\n"
                         "  GET /Batches/085179c3...  → 200 (batch exists)\n"
                         "  GET /Batches/085179c3.../rows → 404 (rows not found)\n"
                         "Reclassified: 15 TCs as CLUSTER_C_PERSISTENCE_SPLIT blocked"),
    ("Expected behaviour","After a successful upload, GET /rows for the same batchId returns the list of uploaded rows "
                          "with their validation status."),
    ("Recommendation",   "Investigate the rows write path. Ensure that on successful upload the rows are "
                         "committed to the rows store atomically with the batch header record. "
                         "If rows are written asynchronously, expose a rows-ready flag on the batch resource "
                         "and document the eventual-consistency window."),
    ("Ceiling unlock",   "Unblocks 15 BLOCKEDs + fixes ~8 FAILs on BATCH-05 (+10.4 pp platform-wide)"),
]
for k, v in fields:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f"{k}: ")
    r.bold = True
    r.font.size = Pt(9)
    p.add_run(v).font.size = Pt(9)

doc.add_paragraph()

# ── D-BATCH-01-VAL ──────────────────────────────────
doc.add_heading("D-BATCH-01-VAL — Upload: File-Type and CSV Header Validation Absent", 2)
fields = [
    ("Defect ID",        "D-BATCH-01-VAL"),
    ("Severity",         "Medium"),
    ("Endpoint",         "POST /api/v1/Batches/card-creation/upload  (BATCH-01)"),
    ("Affected TCs",     "~8 FAILs: unsupported file types accepted, missing required CSV headers not rejected"),
    ("Root Cause",       "The upload handler does not validate: (a) file content-type (PDF, TXT, XLSX accepted alongside CSV); "
                         "(b) required CSV column headers (FirstName, LastName, Phone, Email, dob, idtype, idnumber). "
                         "Files with wrong types or missing required headers return 200 with a batchId."),
    ("Expected behaviour","Non-CSV MIME types (application/pdf, text/plain, application/vnd.openxmlformats-officedocument.spreadsheetml.sheet) "
                          "→ 400 Unsupported Media Type.\n"
                          "CSV missing any required header column → 400 with field-level error identifying the missing column.\n"
                          "Duplicate header columns → 400."),
    ("Recommendation",   "Add a CSV parser step in the upload handler before persisting:\n"
                         "  1. Reject non-CSV MIME types.\n"
                         "  2. Parse the first row as headers.\n"
                         "  3. Assert all required columns are present (case-insensitive).\n"
                         "  4. Assert no duplicate columns.\n"
                         "  5. Return 400 with a structured error identifying the exact violation."),
    ("Ceiling unlock",   "Fixes ~8 FAILs on BATCH-01 (+3.6 pp platform-wide)"),
]
for k, v in fields:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f"{k}: ")
    r.bold = True
    r.font.size = Pt(9)
    p.add_run(v).font.size = Pt(9)

doc.add_paragraph()

# ── D-BATCH-01-PROD ──────────────────────────────────
doc.add_heading("D-BATCH-01-PROD — Upload: Product ID and Row-Count Constraints Not Enforced", 2)
fields = [
    ("Defect ID",        "D-BATCH-01-PROD"),
    ("Severity",         "Medium"),
    ("Endpoint",         "POST /api/v1/Batches/card-creation/upload  (BATCH-01)"),
    ("Affected TCs",     "~4 FAILs"),
    ("Root Cause",       "Requests with an invalid/unknown productId and requests with row counts exceeding the 50-row limit "
                         "are accepted with HTTP 200. No constraint is enforced."),
    ("Expected behaviour","productId not registered for the affiliate → 400 or 422 with field error.\n"
                          "Row count > 50 → 400 with a message indicating the row limit."),
    ("Recommendation",   "Add productId existence and affiliate-eligibility validation.\n"
                         "Add a row-count check: reject uploads where recordsReceived > 50 before persisting."),
    ("Ceiling unlock",   "Fixes ~4 FAILs on BATCH-01 (+1.8 pp platform-wide)"),
]
for k, v in fields:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f"{k}: ")
    r.bold = True
    r.font.size = Pt(9)
    p.add_run(v).font.size = Pt(9)

doc.add_paragraph()

# ── D-BATCH-06-STATE ──────────────────────────────────
doc.add_heading("D-BATCH-06-STATE — Results Download: COMPLETED State Batches Not Consistently Available", 2)
fields = [
    ("Defect ID",        "D-BATCH-06-STATE"),
    ("Severity",         "Medium"),
    ("Endpoint",         "GET /api/v1/Batches/{batchId}/results/download  (BATCH-06)"),
    ("Affected TCs",     "11 FAILs; 2 BLOCKEDs; 19 PASSes (using backend-seeded COMPLETED batch)"),
    ("Root Cause",       "BATCH-06 requires a batch in COMPLETED state. The test environment only provides one "
                         "COMPLETED batch ID (ef57c562-4a98-4c46-b8ec-13e36a1a3ebe) as a seed. "
                         "Scenarios that test error paths (wrong batchId, wrong state) fail because the "
                         "freshly minted batch is never in COMPLETED state."),
    ("Expected behaviour","GET /results/download for a COMPLETED batchId → 200 with downloadUrl (containing token).\n"
                          "GET /results/download for a non-COMPLETED batchId → 409 or 422."),
    ("Recommendation",   "Once D-BATCH-02-STATE / D-BATCH-03-STATE are fixed, the full pipeline "
                         "(upload → validate → submit → process → complete) should produce COMPLETED batches. "
                         "Short-term: provide at least 2 COMPLETED batch IDs from different affiliates "
                         "(see Provisions section)."),
    ("Ceiling unlock",   "Fixes ~6 FAILs on BATCH-06 (+2.7 pp platform-wide) once pipeline is reachable"),
]
for k, v in fields:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f"{k}: ")
    r.bold = True
    r.font.size = Pt(9)
    p.add_run(v).font.size = Pt(9)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════
# 4. FIX SEQUENCE & CEILING PROJECTION
# ═══════════════════════════════════════════════════════
doc.add_heading("4. Fix Sequence & Ceiling Projection", 1)

doc.add_paragraph(
    "Executable TCs: 195 (222 total − 27 structural BLOCKEDs). "
    "Ceiling shown is the projected pass rate after each fix, assuming no regressions."
)

fix_tbl = doc.add_table(rows=1, cols=5)
fix_tbl.style = "Table Grid"
fix_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(fix_tbl, ["Step", "Fix", "FAILs Resolved", "Projected PASSes", "Projected Rate"])

fix_rows = [
    ("0 — Baseline",         "No changes",                    "—",  "85",  "38.3% (85/222)"),
    ("1",  "D-BATCH-AUTH-1: Register auth middleware",         "35", "120", "54.1% (120/222)"),
    ("2",  "D-BATCH-07-CRASH: Implement token handler",        "26", "146", "65.8% (146/222)"),
    ("3",  "D-BATCH-02-STATE: Remove auto-advance race",       "22", "168", "75.7% (168/222)"),
    ("4",  "D-BATCH-03-STATE: Submit requires VALIDATED",      "17", "180", "81.1% (180/222 — ~81% ceiling with current BLOCKEDs)"),
    ("5",  "D-BATCH-05-ROWS: Fix rows persistence split",      "8+15B","188+15", "~88% (188+/222)"),
    ("6",  "D-BATCH-01-VAL: CSV/file-type validation",         "8",  "196", "~88.3%"),
    ("7",  "D-BATCH-01-PROD: ProductId/size constraints",      "4",  "200", "~90.1%"),
    ("8",  "D-BATCH-06-STATE: COMPLETED pipeline fix",         "6",  "206", "~92.8%"),
]
for row_data in fix_rows:
    add_table_row(fix_tbl, list(row_data))

doc.add_paragraph()

# ═══════════════════════════════════════════════════════
# 5. BACKEND PROVISIONS REQUIRED
# ═══════════════════════════════════════════════════════
doc.add_heading("5. Backend Provisions Required", 1)

doc.add_paragraph(
    "The following provisioned IDs / data are required before the next harness run. "
    "All IDs should be stable (not cleaned up between runs) unless marked 'per-run'."
)

prov_tbl = doc.add_table(rows=1, cols=5)
prov_tbl.style = "Table Grid"
prov_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(prov_tbl, ["#", "Resource", "Required Value / Format", "Used By", "Notes"])

provisions = [
    ("P-01", "AFFILIATE_ID",
     "AFF-{hex32} format (e.g. AFF-9F6EDBBE20DD4C6B97D0B720676506E1)",
     "All endpoints",
     "Currently: AFF-9F6EDBBE20DD4C6B97D0B720676506E1 — confirm still valid"),

    ("P-02", "UPLOADED_BATCH_ID",
     "UUID of a batch guaranteed to be in UPLOADED state at test start",
     "BATCH-02 (validate)",
     "Required once D-BATCH-02-STATE is fixed. Batch must remain in UPLOADED until validate is called. "
     "Provide via /api/v1/Batches/card-creation/upload with auth; or freeze a test-fixture batch."),

    ("P-03", "VALIDATED_BATCH_ID",
     "UUID of a batch in VALIDATED state (post-validate, pre-submit)",
     "BATCH-03 (submit)",
     "Required once D-BATCH-02-STATE is fixed so BATCH-03 can test submit happy path."),

    ("P-04", "PROCESSING_BATCH_ID",
     "UUID of a batch in PROCESSING state",
     "BATCH-02, BATCH-03 (negative: wrong-state scenarios)",
     "Currently: 952480b6-61d2-4299-a6ca-430dce7a316c — confirm still valid"),

    ("P-05", "COMPLETED_BATCH_ID (×2)",
     "Two UUIDs for batches in COMPLETED state owned by the seeded affiliate",
     "BATCH-06, BATCH-07",
     "Currently 1 provided: ef57c562-4a98-4c46-b8ec-13e36a1a3ebe. "
     "Need a second COMPLETED batch (from a different upload) to test scope-isolation scenarios."),

    ("P-06", "DOWNLOAD_TOKEN",
     "A valid, non-expired token string from GET /Batches/{completedId}/results/download",
     "BATCH-07",
     "Once D-BATCH-07-CRASH is fixed. Token format should be documented (UUID? hex32? JWT?). "
     "Current captured token: 6483757b22984a2884995a9ec1a6fe10 — likely expired."),

    ("P-07", "FOREIGN_AFFILIATE_COMPLETED_ID",
     "UUID of a COMPLETED batch owned by a *different* affiliate (cross-scope test)",
     "BATCH-06, BATCH-07",
     "Required for foreign_affiliate_token_rejected and scope-isolation tests."),
]
for row_data in provisions:
    add_table_row(prov_tbl, list(row_data))

doc.add_paragraph()

# ═══════════════════════════════════════════════════════
# 6. STRUCTURAL BLOCKEDS
# ═══════════════════════════════════════════════════════
doc.add_heading("6. Structural BLOCKEDs (Irreducible — No Backend Fix Required)", 1)

doc.add_paragraph(
    "The following 27 BLOCKEDs cannot be unblocked by the harness alone. "
    "They are documented here so the test count is not misread as a defect."
)

blk_tbl = doc.add_table(rows=1, cols=4)
blk_tbl.style = "Table Grid"
blk_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(blk_tbl, ["Category", "Count", "Endpoints", "Resolution Path"])

blocked_rows = [
    ("Cluster-C Persistence Split (BATCH-05 rows 404)",
     "15",
     "BATCH-05",
     "Unblocked by D-BATCH-05-ROWS fix; not a harness issue"),
    ("DB-verify-only (need database read to confirm)",
     "6",
     "BATCH-01, BATCH-02, BATCH-03",
     "Requires read-only verification endpoints or DB inspection; out of HTTP-only harness scope"),
    ("Rate-flood (would trigger DoS alarms)",
     "5",
     "BATCH-01",
     "Intentionally skipped; test environment does not support rate/flood testing"),
    ("Unrecognised scenario name",
     "1",
     "BATCH-07",
     "scenario: batch_status_unchanged_after_token_download — add classifier handler in harness"),
]
for row_data in blocked_rows:
    add_table_row(blk_tbl, list(row_data))

doc.add_paragraph()

# ═══════════════════════════════════════════════════════
# 7. SWAGGER CONTRACT ASKS
# ═══════════════════════════════════════════════════════
doc.add_heading("7. Swagger Contract Asks", 1)

doc.add_paragraph(
    "The following additions to the Swagger / OpenAPI spec are requested to improve "
    "signal quality and enable harness contract validation."
)

sw_items = [
    ("SW-01", "Token format documentation",
     "The token returned by BATCH-06 (downloadUrl path segment) has no documented format in swagger. "
     "Add: format (UUID, hex32, JWT?), length, expiry duration, one-time-use policy."),
    ("SW-02", "Batch state machine enum",
     "Add a BatchStatus enum to the swagger schema: UPLOADED, VALIDATING, VALIDATED, "
     "VALIDATION_FAILED, PROCESSING, COMPLETED, FAILED. "
     "Document which state transitions each endpoint triggers."),
    ("SW-03", "Required fields on upload request",
     "The upload request body schema has no 'required' array. "
     "Add: requestContext, productId, file, file.fileName, file.contentType, file.fileBase64."),
    ("SW-04", "Row-count constraint",
     "Document the 50-row maximum on the upload endpoint. "
     "Add x-constraint or pattern to fileBase64, or add a recordsReceived maximum."),
    ("SW-05", "Error response schemas",
     "Batch endpoints return 4xx/5xx without a documented response schema in swagger. "
     "Add 400, 401, 403, 404, 409, 500 response schemas with the standard error envelope."),
]
for sw_id, title, detail in sw_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run(f"{sw_id} — {title}: ")
    r.bold = True
    r.font.size = Pt(9)
    p.add_run(detail).font.size = Pt(9)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════
# 8. APPENDIX — RUN HISTORY
# ═══════════════════════════════════════════════════════
doc.add_heading("8. Appendix — Run History", 1)

run_tbl = doc.add_table(rows=1, cols=5)
run_tbl.style = "Table Grid"
run_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_row(run_tbl, ["Run", "Date", "TCs", "Result", "Key Delta"])

run_history = [
    ("Run 1",  "2026-05-01", "177", "102P/70F/5B (57.6%)",  "Baseline — healthiest service of the day"),
    ("Run 2",  "2026-05-06", "177", "111P/61F/5B (62.7%)",  "+5pp; BATCH-03/04 pack expanded"),
    ("Run 3",  "2026-05-08", "188", "116P/65F/7B (61.7%)",  "+11 BATCH-03/04 TCs; restored silent-accept"),
    ("Run 4",  "2026-05-12", "177", "78P/104F/40B (35.1%)", "BATCH-07 BLOCKEDs exposed; pre-context start"),
    ("Run 5",  "2026-05-18", "222", "88P/106F/28B (39.6%)", "BATCH-07 synthetic entry wired; token probe added"),
    ("Run 6",  "2026-05-18", "222", "86P/108F/28B (38.7%)", "BATCH-04 regression from token-handler gate fix"),
    ("Run 7",  "2026-05-18", "222", "85P/110F/27B (38.3%)", "endpoint gate fix shipped; 0 misfires confirmed"),
    ("Run 8",  "2026-05-18", "222", "85P/110F/27B (38.3%)", "Clean baseline — current report"),
]
for row_data in run_history:
    add_table_row(run_tbl, list(row_data))

doc.add_paragraph()

# ─── save ───────────────────────────────────────────
os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
